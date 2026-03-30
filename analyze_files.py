import pandas as pd
from openpyxl import load_workbook
from docx import Document
import json

def analyze_excel(filepath):
    """Analyze the Excel scanner file"""
    print("=" * 80)
    print("EXCEL FILE ANALYSIS")
    print("=" * 80)

    # Load workbook to see sheet names and formulas
    wb = load_workbook(filepath, data_only=False)
    print(f"\nSheet Names: {wb.sheetnames}")

    # Read each sheet
    excel_data = {}
    for sheet_name in wb.sheetnames:
        print(f"\n{'='*40}")
        print(f"SHEET: {sheet_name}")
        print(f"{'='*40}")

        # Read with pandas for data
        df = pd.read_excel(filepath, sheet_name=sheet_name)
        print(f"Shape: {df.shape}")
        print(f"Columns: {list(df.columns)}")
        print("\nFirst few rows:")
        print(df.head())

        # Get formulas from the sheet (using openpyxl)
        ws = wb[sheet_name]
        formulas = []
        for row in ws.iter_rows(max_row=min(20, ws.max_row)):
            for cell in row:
                if cell.data_type == 'f' and cell.value:
                    formulas.append({
                        'cell': cell.coordinate,
                        'formula': cell.value,
                        'value': cell.internal_value
                    })

        if formulas:
            print("\nSample Formulas found:")
            for f in formulas[:5]:
                print(f"  {f['cell']}: {f['formula']}")

        excel_data[sheet_name] = {
            'columns': list(df.columns),
            'shape': df.shape,
            'sample_data': df.head().to_dict()
        }

    wb.close()
    return excel_data

def analyze_word(filepath):
    """Analyze the strategy document"""
    print("\n" + "=" * 80)
    print("STRATEGY DOCUMENT ANALYSIS")
    print("=" * 80)

    doc = Document(filepath)
    full_text = []

    print("\nDocument Structure:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            full_text.append(para.text)
            if para.style and para.style.name.startswith('Heading'):
                print(f"  [{para.style.name}] {para.text[:80]}...")
            elif len(para.text) < 100 and i < 30:
                print(f"  {para.text}")

    print(f"\nTotal paragraphs: {len([p for p in doc.paragraphs if p.text.strip()])}")
    print(f"Total characters: {sum(len(p.text) for p in doc.paragraphs)}")

    # Try to extract tables
    print("\n" + "=" * 40)
    print("TABLES IN DOCUMENT:")
    print("=" * 40)
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}: {len(table.rows)} rows x {len(table.columns)} columns")
        for row_idx, row in enumerate(table.rows[:min(5, len(table.rows))]):
            row_data = [cell.text.strip() for cell in row.cells]
            print(f"  Row {row_idx+1}: {row_data}")

    return {
        'paragraphs': len([p for p in doc.paragraphs if p.text.strip()]),
        'tables': len(doc.tables),
        'full_text': full_text[:50]  # First 50 paragraphs
    }

# Main execution
if __name__ == "__main__":
    excel_path = r"C:\Users\bhara\OneDrive\Desktop\Trading\DMA_CAR_Enhanced_500.xlsx"
    word_path = r"C:\Users\bhara\OneDrive\Desktop\Trading\DMA_CAR_Strategy_Document_v2.docx"

    print("ANALYZING YOUR TRADING FILES")
    print("=" * 80)

    try:
        excel_analysis = analyze_excel(excel_path)
    except Exception as e:
        print(f"Excel analysis error: {e}")
        excel_analysis = {"error": str(e)}

    try:
        word_analysis = analyze_word(word_path)
    except Exception as e:
        print(f"Word analysis error: {e}")
        word_analysis = {"error": str(e)}

    # Save analysis to JSON
    output = {
        "excel_analysis": excel_analysis,
        "word_analysis": word_analysis
    }

    with open("trading_analysis.json", "w", encoding="utf-8") as f:
        json.dump(output, f, indent=2, default=str)

    print("\n" + "=" * 80)
    print("ANALYSIS COMPLETE")
    print("=" * 80)
    print("Results saved to: trading_analysis.json")
